#!/usr/bin/env python3
"""
build_wps_projection.py — 阶段五：将阶段四三章真实 final.md 组合成 WPS 投影文件

输出目录：/data/webnovel-lab/demo_output/phase5_wps_projection/
输出文件：
  1. price_tag_life_ch001.md
  2. price_tag_life_ch002.md
  3. price_tag_life_ch003.md
  4. price_tag_life_volume_001.md
  5. price_tag_life_volume_001.docx
  6. projection_manifest.json
"""

import os
import sys
import json
import hashlib
from pathlib import Path
from datetime import datetime


CODE_ROOT = Path(os.environ.get("WEBNOVEL_CODE_ROOT", "/opt/webnovel-hermes-wps"))
DATA_ROOT = Path(os.environ.get("WEBNOVEL_DATA_ROOT", "/data/webnovel-lab"))


def main():
    import argparse
    parser = argparse.ArgumentParser(description="阶段五 WPS 投影文件生成")
    parser.add_argument("--project", default="price_tag_life", help="项目名")
    args = parser.parse_args()

    project = args.project
    ch_titles = {
        1: "价格初现",
        2: "误判代价",
        3: "第一次主动选择",
    }

    # ── 输入来源 ──
    ch1_path = DATA_ROOT / "demo_output" / "phase4b_real_run" / "chapter_001" / "final.md"
    ch2_path = DATA_ROOT / "demo_output" / "phase4c_real_run" / "chapter_002" / "final.md"
    ch3_path = DATA_ROOT / "demo_output" / "phase4c_real_run" / "chapter_003" / "final.md"

    for ch, p in [(1, ch1_path), (2, ch2_path), (3, ch3_path)]:
        if not p.exists():
            print(f"❌ chapter_{ch:03d} final.md 缺失: {p}")
            sys.exit(1)
        print(f"  ✅ Chapter {ch}: {p} ({p.stat().st_size} bytes)")

    # ── 读取正文 ──
    ch_texts = {}
    for ch, p in [(1, ch1_path), (2, ch2_path), (3, ch3_path)]:
        raw = p.read_text(encoding="utf-8")
        # 过滤掉 canon_check 自检段落（从 --- 之后或 canon_check 标记）
        lines = raw.split("\n")
        body_lines = []
        in_canon_check = False
        for line in lines:
            if line.strip().startswith("**canon_check:**") or line.strip().startswith("canon_check:"):
                in_canon_check = True
                continue
            if in_canon_check and line.strip().startswith("---"):
                in_canon_check = False
                continue
            if in_canon_check:
                continue
            body_lines.append(line)
        clean = "\n".join(body_lines).strip()
        ch_texts[ch] = clean
        # 估算中文字数
        cn_count = sum(1 for c in clean if '\u4e00' <= c <= '\u9fff')
        print(f"  📝 Chapter {ch}: {cn_count} 中文字 (正文部分)")

    # ── 输出目录 ──
    out_dir = DATA_ROOT / "demo_output" / "phase5_wps_projection"
    out_dir.mkdir(parents=True, exist_ok=True)

    # ── 1. 单章 Markdown ──
    for ch in [1, 2, 3]:
        title = ch_titles[ch]
        md_path = out_dir / f"{project}_ch{ch:03d}.md"
        content = f"# 第{['一','二','三'][ch-1]}章 {title}\n\n{ch_texts[ch]}"
        md_path.write_text(content, encoding="utf-8")
        print(f"  ✅ {md_path.name} ({md_path.stat().st_size} bytes)")

    # ── 2. 合集 Markdown ──
    volume_md = out_dir / f"{project}_volume_001.md"
    vol_lines = [
        "# 人生价格标签",
        "",
        "## 第一卷：价格初现",
        "",
    ]
    for ch in [1, 2, 3]:
        title = ch_titles[ch]
        vol_lines.append(f"### 第{['一','二','三'][ch-1]}章 {title}")
        vol_lines.append("")
        vol_lines.append(ch_texts[ch])
        vol_lines.append("")
        vol_lines.append("---")
        vol_lines.append("")

    volume_md.write_text("\n".join(vol_lines), encoding="utf-8")
    print(f"  ✅ {volume_md.name} ({volume_md.stat().st_size} bytes)")

    # ── 3. DOCX ──
    docx_path = out_dir / f"{project}_volume_001.docx"
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0.74)

        # cover
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("人生价格标签")
        run.font.size = Pt(28)
        run.font.name = "SimHei"

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run("第一卷：价格初现")
        run2.font.size = Pt(16)
        run2.font.name = "SimHei"

        doc.add_page_break()

        for ch in [1, 2, 3]:
            title = ch_titles[ch]
            heading = doc.add_heading(f"第{['一','二','三'][ch-1]}章 {title}", level=1)
            for line in ch_texts[ch].split("\n"):
                if line.strip().startswith("#"):
                    doc.add_heading(line.strip("# ").strip(), level=2)
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("")

        doc.save(str(docx_path))
        docx_size = docx_path.stat().st_size
        print(f"  ✅ {docx_path.name} ({docx_size} bytes)")
    except ImportError:
        print("  ⚠️ python-docx 不可用，跳过 DOCX 生成")
        docx_size = 0
    except Exception as e:
        print(f"  ❌ DOCX 生成失败: {e}")
        docx_size = 0

    # ── 4. sha256 (仅 DOCX) ──
    sha256_val = ""
    if docx_path.exists() and docx_size > 0:
        sha256_val = hashlib.sha256(docx_path.read_bytes()).hexdigest()
        sha_path = out_dir / f"{project}_volume_001.docx.sha256"
        sha_path.write_text(sha256_val, encoding="utf-8")
        print(f"  ✅ {sha_path.name}: {sha256_val[:16]}...")

    # ── 5. projection_manifest.json ──
    chapters_manifest = []
    for ch in [1, 2, 3]:
        cn_count = sum(1 for c in ch_texts[ch] if '\u4e00' <= c <= '\u9fff')
        chapters_manifest.append({
            "chapter_id": f"chapter_{ch:03d}",
            "title": ch_titles[ch],
            "source": str(DATA_ROOT / ("demo_output/phase4b_real_run" if ch == 1 else "demo_output/phase4c_real_run") / f"chapter_{ch:03d}" / "final.md"),
            "markdown": str(out_dir / f"{project}_ch{ch:03d}.md"),
            "char_count": cn_count,
        })

    manifest = {
        "project": project,
        "title": "人生价格标签",
        "volume": "volume_001",
        "chapters": chapters_manifest,
        "combined_markdown": str(volume_md),
        "docx": str(docx_path) if docx_size > 0 else "",
        "docx_sha256": sha256_val,
        "docx_size_bytes": docx_size,
        "created_at": datetime.now().isoformat(),
        "source_phase": "phase4_real",
    }

    manifest_path = out_dir / "projection_manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✅ projection_manifest.json")
    print(f"     章节数: {len(chapters_manifest)}")
    print(f"     总字数: {sum(c['char_count'] for c in chapters_manifest)} 中文字")

    # ── 输出摘要 ──
    print(f"\n{'='*50}")
    print(f"  WPS 投影生成完成 ✅")
    print(f"  输出目录: {out_dir}")
    print(f"  {'='*50}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
