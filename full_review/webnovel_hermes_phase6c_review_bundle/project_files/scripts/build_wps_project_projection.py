#!/usr/bin/env python3
"""
build_wps_project_projection.py — Phase 6C: WPS 项目化本地投影生成

输入来源：phase5_wps_projection / phase4b+4c real 产出 / .story-system
输出目录：/data/webnovel-lab/demo_output/phase6c_wps_project_projection/

不修改正文、不修改 .story-system、不调用 DeepSeek real。
"""

import os, sys, json, yaml, hashlib
from pathlib import Path
from datetime import datetime

PROJECT_ROOT = Path("/opt/webnovel-hermes-wps")
DATA_ROOT = Path("/data/webnovel-lab")
OUTPUT_ROOT = DATA_ROOT / "demo_output" / "phase6c_wps_project_projection"

PROJECT_ID = "price_tag_life"
TODAY = datetime.now().strftime("%Y-%m-%d")
VERSION = "v001"

# 输入文件
VOLUME_MD    = DATA_ROOT / "demo_output" / "phase5_wps_projection" / f"{PROJECT_ID}_volume_001.md"
CH_FINALS = {
    1: DATA_ROOT / "demo_output" / "phase4b_real_run" / "chapter_001" / "final.md",
    2: DATA_ROOT / "demo_output" / "phase4c_real_run" / "chapter_002" / "final.md",
    3: DATA_ROOT / "demo_output" / "phase4c_real_run" / "chapter_003" / "final.md",
}
RUNTIME_CANON_PATH  = DATA_ROOT / "demo_output" / "phase4c_real_run" / "runtime_canon_real_ch001_to_ch003.yaml"
STORY_SYSTEM        = DATA_ROOT / "workspace" / "novels" / PROJECT_ID / ".story-system"
MASTER_SETTING_PATH = STORY_SYSTEM / "MASTER_SETTING.yaml"
CANON_PATTERNS_PATH = STORY_SYSTEM / "canon_patterns.yaml"
DEBTS_PATH          = STORY_SYSTEM / "reader_debts.yaml"

CH_TITLES = {1: "价格初现", 2: "误判代价", 3: "第一次主动选择"}
FOLDER_MAP = {"volume_current": "00_发布稿", "story_bible": "01_设定资料", "wiki": "02_小说Wiki", "archive": "99_归档版本"}


def read_t(path):
    try: return path.read_text(encoding="utf-8")
    except: return ""

def write_f(path, content):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"  ✅ {path.relative_to(OUTPUT_ROOT)}")

def sha256(path):
    if path.exists() and path.stat().st_size > 0:
        return hashlib.sha256(path.read_bytes()).hexdigest()
    return ""

def build_docx(md_path, docx_path):
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "SimSun"
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0.74)
        text = read_t(md_path)
        for line in text.split("\n"):
            s = line.strip()
            if s.startswith("# "): doc.add_heading(s[2:], level=1)
            elif s.startswith("## "): doc.add_heading(s[3:], level=2)
            elif s.startswith("### "): doc.add_heading(s[4:], level=3)
            elif s: doc.add_paragraph(s)
            else: doc.add_paragraph("")
        doc.save(str(docx_path))
    except ImportError:
        print(f"  ⚠️ python-docx 不可用，空文件")
        docx_path.write_bytes(b"")
    except Exception as e:
        print(f"  ⚠️ DOCX 异常: {e}")
        docx_path.write_bytes(b"")


def main():
    # ════════ 1. 发布稿 ════════
    pub_dir = OUTPUT_ROOT / "publish"
    pub_md = pub_dir / "人生价格标签_第一卷_当前版.md"
    pub_docx = pub_dir / "人生价格标签_第一卷_当前版.docx"

    if VOLUME_MD.exists():
        text = read_t(VOLUME_MD)
    else:
        parts = ["# 人生价格标签\n", "## 第一卷：价格初现\n"]
        for ch in [1, 2, 3]:
            parts.append(f"\n### 第{['一','二','三'][ch-1]}章 {CH_TITLES[ch]}\n")
            parts.append(read_t(CH_FINALS[ch]))
            parts.append("\n---\n")
        text = "".join(parts)

    write_f(pub_md, text)
    build_docx(pub_md, pub_docx)
    write_f(pub_dir / f"{pub_docx.stem}.docx.sha256", sha256(pub_docx))
    print(f"  发布稿: {pub_md.stat().st_size} chars, DOCX: {pub_docx.stat().st_size} bytes")

    # ════════ 2. Story Bible ════════
    bible_dir = OUTPUT_ROOT / "story_bible"
    bible_md = bible_dir / "人生价格标签_Story_Bible_当前版.md"
    bible_docx = bible_dir / "人生价格标签_Story_Bible_当前版.docx"

    master = yaml.safe_load(read_t(MASTER_SETTING_PATH)) or {}
    p = master.get("project", {})
    chars = master.get("characters", [])
    w = master.get("world", {})
    o = master.get("outline", {})
    cp = yaml.safe_load(read_t(CANON_PATTERNS_PATH)) or {}
    rc = yaml.safe_load(read_t(RUNTIME_CANON_PATH)) or {}

    lines = [
        "# 人生价格标签 — Story Bible",
        f"> 更新: {TODAY} | 来源: .story-system/MASTER_SETTING.yaml",
        "", "## 1. 书名", p.get("name", PROJECT_ID),
        "", "## 2. 类型", p.get("genre", ""),
        "", "## 3. 核心脑洞", p.get("core_idea", ""),
        "", "## 4. 目标读者", p.get("target_reader", ""),
        "", "## 5. 风格偏好", p.get("style_preference", ""),
        "", "## 6. 主角设定",
    ]
    for c in chars:
        lines.append(f"- **{c.get('name','?')}**（{c.get('age','?')}岁，{c.get('identity','?')}）")
        for k in ["personality", "motivation", "arc"]:
            if c.get(k): lines.append(f"  {k}: {c[k]}")
        if c.get("initial_price_tag"): lines.append(f"  初始标签: {c['initial_price_tag']}")
        if c.get("price_tag_meaning"): lines.append(f"  标签含义: {c['price_tag_meaning']}")

    lines += ["", "## 7. 世界观背景", w.get("background", "")]
    lines += ["", "## 8. 能力规则"]
    for r in w.get("rules", []): lines.append(f"- {r}")
    lines += ["", "## 9. 禁止设定"]
    for fb in cp.get("forbidden_patterns", []): lines.append(f"- 禁止: {fb}")
    lines += ["", "## 10. 第一卷方向", f"当前弧: {o.get('current_arc','')}", f"总章节: {o.get('total_chapters','')}"]
    for ka in o.get("key_arcs", []): lines.append(f"- {ka}")
    lines += ["", "## 11. 前三章摘要"]
    for ch in [1, 2, 3]:
        ch_ev = [e["event"] for e in rc.get("confirmed_events",[]) if e.get("chapter")==ch]
        lines.append(f"\n### 第{['一','二','三'][ch-1]}章")
        for ev in ch_ev[:5]: lines.append(f"- {ev}")

    write_f(bible_md, "\n".join(lines))
    build_docx(bible_md, bible_docx)
    write_f(bible_dir / f"{bible_docx.stem}.docx.sha256", sha256(bible_docx))
    print(f"  Story Bible: {bible_md.stat().st_size} chars, DOCX: {bible_docx.stat().st_size} bytes")

    # ════════ 3. 小说 Wiki ════════
    wiki_dir = OUTPUT_ROOT / "wiki"
    wiki_md = wiki_dir / "人生价格标签_小说Wiki_当前版.md"
    wiki_docx = wiki_dir / "人生价格标签_小说Wiki_当前版.docx"

    lines = [
        "# 人生价格标签 — 小说Wiki",
        f"> 更新: {TODAY} | 来源: runtime_canon + .story-system",
        "", "## 1. 已发生事件",
    ]
    for ch in [1, 2, 3]:
        ch_ev = [e["event"] for e in rc.get("confirmed_events",[]) if e.get("chapter")==ch]
        lines.append(f"\n### 第{['一','二','三'][ch-1]}章")
        for ev in ch_ev: lines.append(f"- {ev}")

    proto = rc.get("protagonist", {})
    lines += ["", "## 2. 角色当前状态"]
    lines.append(f"- **林砚**: {proto.get('name','?')}，{proto.get('identity','?')}")
    for k, v in proto.get("status",{}).items(): lines.append(f"  - {k}: {v}")
    for c in rc.get("characters",[]):
        if isinstance(c, dict): lines.append(f"- **{c.get('name',c.get('id','?'))}**: {c.get('current_state','')}")

    s = proto.get("status", {})
    lines += ["", "## 3. 能力理解进度",
        f"- 觉醒: {s.get('ability_awakened','?')}", f"- 理解标签: {s.get('understands_label','?')}",
        "", "## 4. 开放线索"]
    for thr in rc.get("open_threads", []):
        lines.append(f"- {thr.get('id','?')}: {thr.get('status','?')}")

    lines += ["", "## 5. 伏笔线索"]
    if DEBTS_PATH.exists():
        debts_data = yaml.safe_load(read_t(DEBTS_PATH))
        if isinstance(debts_data, dict):
            debts_data = debts_data.get("debts", [])
        for dbt in debts_data or []:
            lines.append(f"- [{dbt.get('status','?')}] {dbt.get('description','?')} (ch{dbt.get('created_chapter','?')})")

    lines += ["", "## 6. 禁止事实"]
    for fb in cp.get("forbidden_patterns",[]): lines.append(f"- 禁止: {fb}")
    lines += ["", "## 7. 章节索引"]
    for ch in [1, 2, 3]: lines.append(f"- 第{['一','二','三'][ch-1]}章「{CH_TITLES[ch]}」✅")

    write_f(wiki_md, "\n".join(lines))
    build_docx(wiki_md, wiki_docx)
    write_f(wiki_dir / f"{wiki_docx.stem}.docx.sha256", sha256(wiki_docx))
    print(f"  Wiki: {wiki_md.stat().st_size} chars, DOCX: {wiki_docx.stat().st_size} bytes")

    # ════════ 4. 章节索引 ════════
    idx_dir = OUTPUT_ROOT / "index"
    write_f(idx_dir / "人生价格标签_章节索引.md", "\n".join([
        "# 章节索引", "", "| 章 | 标题 | 状态 |",
        "|---|---|---|",
        "| 1 | 价格初现 | ✅ 完成 |",
        "| 2 | 误判代价 | ✅ 完成 |",
        "| 3 | 第一次主动选择 | ✅ 完成 |",
        "| 4+ | — | ⏸ 待生产 |", "",
    ]))
    write_f(idx_dir / "人生价格标签_角色状态表.md", "\n".join([
        "# 角色状态表", "",
        "| 角色 | 身份 | 状态 | 标签 |",
        "|---|---|---|---|",
        "| 林砚 | 外卖员 | 觉醒中 | 归零边缘 |",
        "| 林父 | 病人 | 病重 | 正在归零 |",
        "| 老人 | 神秘 | 已知能力 | 异常高 |",
        "| 客户 | 上班族 | 财务危机 | 快速下跌 |", "",
    ]))
    write_f(idx_dir / "人生价格标签_伏笔线索表.md", "\n".join([
        "# 伏笔线索表", "",
        "| 线索 | 类型 | 状态 |",
        "|---|---|---|",
        "| 父亲归零 | 主线 | open |",
        "| 老人身份 | 支线 | open |",
        "| 标签含义 | 核心 | open |",
        "| 能力代价 | 设定 | confirmed |", "",
    ]))

    # ════════ 5. 归档 ════════
    arch_dir = OUTPUT_ROOT / "archive"
    arch_dir.mkdir(parents=True, exist_ok=True)
    arch_name = f"{TODAY}_v001_第一卷前三章"
    arch_docx = arch_dir / f"{arch_name}.docx"
    if pub_docx.exists() and pub_docx.stat().st_size > 0:
        arch_docx.write_bytes(pub_docx.read_bytes())
    else:
        build_docx(pub_md, arch_docx)
    write_f(arch_dir / f"{arch_name}.docx.sha256", sha256(arch_docx))

    # ════════ 6. manifest ════════
    docs_list = []
    for dk, title, path in [
        ("volume_current", "人生价格标签_第一卷_当前版", pub_docx),
        ("story_bible", "人生价格标签_Story_Bible_当前版", bible_docx),
        ("wiki", "人生价格标签_小说Wiki_当前版", wiki_docx),
        ("archive", arch_name, arch_docx),
    ]:
        docs_list.append({"doc_key": dk, "title": title, "local_path": str(path),
                          "target_folder": FOLDER_MAP[dk], "format": "docx", "sha256": sha256(path) if path.exists() else ""})
    manifest = {"project_id": PROJECT_ID, "title": "人生价格标签", "source_truth": ".story-system",
                "projection_type": "wps_project", "created_at": datetime.now().isoformat(),
                "project_folder": "小说/人生价格标签",
                "folders": ["00_发布稿","01_设定资料","02_小说Wiki","03_章节索引","99_归档版本"],
                "documents": docs_list, "version": VERSION, "redacted": True}
    write_f(OUTPUT_ROOT / "wps_project_manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))

    structure = {"project_folder": "小说/人生价格标签", "folders": [
        {"key": "publish", "name": "00_发布稿"},
        {"key": "story_bible", "name": "01_设定资料"},
        {"key": "wiki", "name": "02_小说Wiki"},
        {"key": "index", "name": "03_章节索引"},
        {"key": "archive", "name": "99_归档版本"},
    ]}
    write_f(OUTPUT_ROOT / "wps_project_structure.yaml", yaml.dump(structure, allow_unicode=True, default_flow_style=False))

    print(f"\n✅ 输出: {OUTPUT_ROOT}")
    print(f"   文件数: {sum(1 for _ in OUTPUT_ROOT.rglob('*') if _.is_file())}")
    return 0

if __name__ == "__main__":
    sys.exit(main())
