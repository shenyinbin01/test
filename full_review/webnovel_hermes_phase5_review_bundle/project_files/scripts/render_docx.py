#!/usr/bin/env python3
"""
render_docx.py — 将小说项目渲染为 DOCX 和 Markdown
依赖: python-docx (如不可用，降级为 Markdown 输出)
"""

import sys
import os
from pathlib import Path


CODE_ROOT = Path(os.environ.get("WEBNOVEL_CODE_ROOT", "/opt/webnovel-hermes-wps"))
DATA_ROOT = Path(os.environ.get("WEBNOVEL_DATA_ROOT", "/data/webnovel-lab"))


def find_demo_project():
    """查找 Demo 项目路径"""
    novels_dir = DATA_ROOT / "workspace" / "novels"
    candidates = ["price_tag_life"]
    for name in candidates:
        project_dir = novels_dir / name
        if project_dir.exists():
            return project_dir, name
    return None, None


def collect_chapters(project_dir):
    """收集已完成章节的最终正文"""
    chapters_dir = project_dir / "manuscript" / "chapters"
    if not chapters_dir.exists():
        return []

    chapters = sorted(chapters_dir.glob("*_final.md"))
    return chapters


def collect_drafts(project_dir):
    """如果没有 final 正文，回退到 drafts"""
    drafts_dir = project_dir / "manuscript" / "drafts"
    if not drafts_dir.exists():
        return []
    return sorted(drafts_dir.glob("*_draft.md"))


def render_markdown(project_dir, output_path):
    """渲染为 Markdown"""
    project_name = project_dir.name
    lines = []

    # 尝试读取 project.yaml
    project_yaml = project_dir / "project.yaml"
    book_title = project_name
    if project_yaml.exists():
        try:
            import yaml
            data = yaml.safe_load(project_yaml.read_text())
            if data and "book_title" in data:
                book_title = data["book_title"]
        except Exception:
            pass

    lines.append(f"# {book_title}")
    lines.append("")
    lines.append(f"> 项目: {project_name}")
    lines.append("")

    chapters = collect_chapters(project_dir)
    if not chapters:
        chapters = collect_drafts(project_dir)

    if not chapters:
        lines.append("*暂无已完成章节*")
    else:
        for ch in chapters:
            content = ch.read_text(encoding="utf-8")
            lines.append(content)
            lines.append("")
            lines.append("---")
            lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"✅ Markdown 已输出: {output_path}")
    return True


def render_docx(project_dir, output_path):
    """渲染为 DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠️ python-docx 不可用，降级为 Markdown 输出")
        md_path = output_path.with_suffix(".md")
        return render_markdown(project_dir, md_path)

    project_name = project_dir.name
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)

    # 读取书名
    project_yaml = project_dir / "project.yaml"
    book_title = project_name
    if project_yaml.exists():
        try:
            import yaml
            data = yaml.safe_load(project_yaml.read_text())
            if data and "book_title" in data:
                book_title = data["book_title"]
        except Exception:
            pass

    # 封面
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(book_title)
    run.font.size = Pt(28)
    run.font.name = "SimHei"

    doc.add_page_break()

    chapters = collect_chapters(project_dir)
    if not chapters:
        chapters = collect_drafts(project_dir)

    if not chapters:
        doc.add_paragraph("暂无已完成章节")
    else:
        for ch in chapters:
            content = ch.read_text(encoding="utf-8")
            lines = content.split("\n")
            for line in lines:
                if line.startswith("# 第") and "章" in line:
                    heading = doc.add_heading(line.strip("# "), level=1)
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("")

    doc.save(str(output_path))
    print(f"✅ DOCX 已输出: {output_path}")
    return True


def main():
    project_dir, project_name = find_demo_project()
    if not project_dir:
        print("⚠️ 未找到 Demo 项目。请先运行 run_demo.py。")
        # 仍然生成占位输出
        exports_dir = DATA_ROOT / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)
        placeholder = exports_dir / "无项目_output.md"
        placeholder.write_text("# 未找到项目\n\n请先运行 python scripts/run_demo.py 创建 Demo 项目。")
        print(f"✅ 占位文件已输出: {placeholder}")
        return

    exports_dir = project_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    # 输出 DOCX
    docx_path = exports_dir / f"{project_name}_main.docx"
    render_docx(project_dir, docx_path)

    # 输出 Markdown
    md_path = exports_dir / f"{project_name}_main.md"
    render_markdown(project_dir, md_path)


if __name__ == "__main__":
    main()
